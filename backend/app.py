from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile
import zipfile
from io import BytesIO
import json
import sqlite3
from datetime import datetime
import re


try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
    print("WARNING: PyPDF2 not found. PDF extraction will not work. Install with 'pip install PyPDF2'")

app = Flask(__name__)
CORS(app, supports_credentials=True)

print("Flask app is initializing...")

class TemplateDatabase:
    """Manages document templates in an SQLite database."""
    def __init__(self, db_path="templates.db"):
        self.db_path = db_path
        self.init_database()

    def init_database(self):
        """Initializes the SQLite database table for templates."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT UNIQUE,
                content BLOB,
                upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                format_info TEXT
            )
        ''')
        conn.commit()
        conn.close()

    def save_template(self, name, content, format_info):
        """Saves or updates a template in the database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        try:
            # Check if template already exists to decide between INSERT and UPDATE
            cursor.execute('SELECT id FROM templates WHERE name = ?', (name,))
            existing_template = cursor.fetchone()

            if existing_template:
                # Update existing template
                cursor.execute('''
                    UPDATE templates
                    SET content = ?, format_info = ?, upload_date = CURRENT_TIMESTAMP
                    WHERE name = ?
                ''', (content, json.dumps(format_info), name))
                print(f"DEBUG: Template '{name}' updated in DB.")
            else:
                # Insert new template
                cursor.execute('''
                    INSERT INTO templates (name, content, format_info)
                    VALUES (?, ?, ?)
                ''', (name, content, json.dumps(format_info)))
                print(f"DEBUG: Template '{name}' inserted into DB.")
            conn.commit()
            return True
        except Exception as e:
            print(f"Error saving template: {e}")
            return False
        finally:
            conn.close()

    def get_templates(self):
        """Retrieves all templates from the database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT name, content, format_info FROM templates ORDER BY upload_date DESC')
        templates = cursor.fetchall()
        conn.close()
        return templates

    def delete_template(self, name):
        """Deletes a specific template from the database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('DELETE FROM templates WHERE name = ?', (name,))
        conn.commit()
        conn.close()
        print(f"DEBUG: Template '{name}' deleted from DB.")

class DocumentGenerator:
    """Handles document generation and formatting."""
    def __init__(self):
        self.template_db = TemplateDatabase()

    def extract_template_format(self, doc_content):
        """Extracts formatting from a DOCX file."""
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(doc_content)
            tmp_path = tmp.name

        try:
            doc = Document(tmp_path)
            format_info = {
                'paragraphs': [],
                'styles': {}, # styles are more complex to extract fully, leaving as placeholder
                'runs': []    # run-level info is stored within paragraphs
            }

            for i, paragraph in enumerate(doc.paragraphs):
                alignment_name = None
                if paragraph.alignment is not None:
                    # Map EnumValue to a string representation
                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                        alignment_name = 'left'
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment_name = 'center'
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        alignment_name = 'right'
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment_name = 'justify'
                    else:
                        alignment_name = 'left' # Default or handle other alignments if needed

                para_format = {
                    'text': paragraph.text, # Store original text for debugging/reference
                    'alignment': alignment_name,
                    'space_before': paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else None,
                    'space_after': paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else None,
                    'line_spacing': paragraph.paragraph_format.line_spacing,
                    'runs': [] # Store run-level info for this paragraph
                }

                for run in paragraph.runs:
                    run_format = {
                        'text': run.text, # Store original run text for reference
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'font_color': str(run.font.color.rgb) if run.font.color else None
                    }
                    para_format['runs'].append(run_format)

                format_info['paragraphs'].append(para_format)

            print(f"DEBUG: Extracted format info for {len(format_info['paragraphs'])} paragraphs.")
            return format_info, doc

        finally:
            os.unlink(tmp_path)

    def generate_content_with_placeholders(self, template_content_bytes, project_desc, lead_firm, jv_firm, date, gemini_api_key, person_name, current_location, person_address, project_title, country):
        """Uses Gemini AI to replace placeholders in template content."""
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(template_content_bytes)
                tmp_path = tmp.name

            doc = Document(tmp_path)
            
            template_full_text = "\n".join([p.text for p in doc.paragraphs])

            
            genai.configure(api_key=gemini_api_key)
            model = genai.GenerativeModel('gemini-1.5-flash')

            prompt = f"""
            You are an AI designed to perform precise text replacement in a document template.
            Your task is to replace specific placeholders in the provided document template.

            **STRICT RULES FOR YOUR RESPONSE:**
            1.  **DO NOT ADD OR REMOVE ANY CHARACTERS OR NEWLINES** other than the actual placeholder replacements.
            2.  **PRESERVE ALL ORIGINAL WHITESPACE, BLANK LINES, AND SPECIAL CHARACTERS (e.g., "") EXACTLY.**
            3.  **ONLY replace the text within the square brackets** that match the specified placeholders.
            4.  **The output MUST be the exact content of the template with ONLY the replacements made.**
            5.  **Do NOT include any introductory phrases, concluding remarks, or explanations.**
            6.  **[PLACE] and [CURRENT_DATE] should be replaced such that they appear on the same line, separated by a comma.**

            **Document Template (copy this text exactly, only replace placeholders):**
            {template_full_text}

            **Replacements to make:**
            - [PROJECT_DESCRIPTION_TEXT] -> {project_desc}
            - [LEAD_FIRM_NAME] -> {lead_firm}
            - [JV_FIRM_NAME] -> {jv_firm if jv_firm else ''}
            - [CURRENT_DATE] -> {date}
            - [PERSON_NAME] -> {person_name if person_name else ''}
            - [PLACE] -> {country if country else current_location}
            - [ASSIGNMENT NAME] -> {project_title if project_title else project_desc}
            - [Company names of the Joint Venture partners/Sub Contractors] -> {jv_firm if jv_firm else ''}
            - [Registered Address of the Joint Venture partners/Sub Contractors] -> {''}
            - [Name of the person in charge of the Lead Firm] -> {person_name if person_name else ''}
            - [Title of the person in charge of the Lead Firm] -> {''}
            - [Role of the Lead Firm] -> {''}
            - [phone number the Lead Firm] -> {''}
            - [Email Address of the Lead Firm] -> {''}
            - [To address] -> {person_address if person_address else ''}

            Example of [PLACE], [Date] replacement: "{country if country else current_location}, {date}"
            """
            response = model.generate_content(prompt)
            # Split the response text by original paragraphs, assuming Gemini perfectly preserves them
            # This is the ideal scenario we are prompting Gemini for.
            return response.text.split('\n')

        except Exception as e:
            print(f"Error generating content with Gemini: {str(e)}")
            raise  # Re-raise the exception after printing
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)  # Ensure temporary file is deleted


    def apply_template_formatting_exact(self, format_info, new_content_lines):
        """
        AThe Document Generation Flask Application has been updated to use the
        extracted person's address for the "[To address]" placeholder.

        Key Changes:
        * The `generate_content_with_placeholders` function now accepts `person_address`
            as an argument.
        * The prompt within `generate_content_with_placeholders` has been modified
            to replace `[To address]` with `{person_address if person_address else ''}`.
        * The call to `generate_content_with_placeholders` in the `generate_document`
            route has been updated to pass the `person_address` parameter.
        """
        new_doc = Document()
        # Create an iterator for the generated content lines
        new_content_iterator = iter(new_content_lines)

        # Mapping for alignment values
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }

        # Keep track of which original template paragraph we are currently trying to match
        original_para_index = 0

        while original_para_index < len(format_info['paragraphs']):
            original_para_format = format_info['paragraphs'][original_para_index]
            original_text = original_para_format['text']

            new_paragraph_text = ""
            try:
                new_paragraph_text = next(new_content_iterator)
                # print(f"DEBUG: Matching original_para_idx={original_para_index} ('{original_text.strip()[:50]}...') with new_text='{new_paragraph_text.strip()[:50]}...'")

            except StopIteration:
              
                print(f"DEBUG: Ran out of generated content lines at original_para_idx={original_para_index}. Stopping.")
                break

            p = new_doc.add_paragraph(new_paragraph_text)

            # Apply paragraph-level formatting from the original template
            if original_para_format['alignment'] in alignment_map:
                p.alignment = alignment_map[original_para_format['alignment']]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT # Default

            if original_para_format['space_before'] is not None:
                p.paragraph_format.space_before = Pt(original_para_format['space_before'])
            if original_para_format['space_after'] is not None:
                p.paragraph_format.space_after = Pt(original_para_format['space_after'])
            if original_para_format['line_spacing'] is not None:
                p.paragraph_format.line_spacing = original_para_format['line_spacing']

            # Apply run-level formatting
            if original_para_format['runs'] and p.runs:
                first_run_format = original_para_format['runs'][0]
                run = p.runs[0]
                if first_run_format.get('bold'): run.bold = True
                if first_run_format.get('italic'): run.italic = True
                if first_run_format.get('underline'): run.underline = True
                if first_run_format.get('font_name'): run.font.name = first_run_format['font_name']
                if first_run_format.get('font_size'): run.font.size = Pt(first_run_format['font_size'])

            original_para_index += 1 # Move to the next original template paragraph

        for remaining_line in new_content_iterator:
            if remaining_line.strip(): # Only add if it contains actual content
                new_doc.add_paragraph(remaining_line)
                print(f"DEBUG: Added remaining line with default format: '{remaining_line.strip()[:50]}...'")

        return new_doc


    def extract_description_with_gemini(self, file_content, file_type, gemini_api_key):
        """Extracts project description, person name, address, project title, and country using Gemini API from various file types."""
        if not gemini_api_key:
            raise ValueError("Gemini API Key not provided for description extraction.")

        genai.configure(api_key=gemini_api_key)

        model = genai.GenerativeModel('gemini-1.5-flash')

        text_content = ""
        if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Handle .docx
            try:
                doc = Document(BytesIO(file_content))
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
            except Exception as e:
                raise ValueError(f"Error reading DOCX file: {e}")
        elif file_type == 'application/pdf':
            # Handle .pdf
            if PdfReader:
                try:
                    reader = PdfReader(BytesIO(file_content))
                    for page in reader.pages:
                        text_content += page.extract_text() + "\n"
                except Exception as e:
                    raise ValueError(f"Error reading PDF file: {e}")
            else:
                raise ValueError("PDF extraction requires PyPDF2. Please install it.")
        elif file_type == 'text/plain':
            # Handle .txt
            try:
                text_content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text_content = file_content.decode('latin-1', errors='ignore') # Fallback for some encodings
        else:
            raise ValueError(f"Unsupported file type for description extraction: {file_type}")

        if not text_content.strip():
            return {"description": "No readable text content found in the uploaded file.", "person_name": "N/A", "person_address": "N/A", "project_title": "N/A", "country": "N/A"}

        # Prompt for project description, person name, address, project title, and country
        prompt = f"""
        Analyze the following text content and extract the following information. If a piece of information is not found, indicate "N/A".

        1.  **Project Description**: A concise summary focusing on key objectives, technologies, and deliverables.
        2.  **Project Title**: A short, formal title for the project.
        3.  **Person Name**: The full name of any person mentioned that appears to be a primary contact or recipient of a letter/document.
        4.  **Person Address**: The full address associated with that person or the project itself. If an address contains multiple lines, keep them.
        5.  **Country**: The country most relevant to the project or the address mentioned.

        Text:
        {text_content}

        Output should be in JSON format like this:
        {{
            "project_description": "...",
            "project_title": "...",
            "person_name": "...",
            "person_address": "...",
            "country": "..."
        }}
        """
        response = model.generate_content(prompt)

        try:
            # Attempt to parse the response as JSON
            extracted_data = json.loads(response.text)
            description = extracted_data.get("project_description", "N/A")
            project_title = extracted_data.get("project_title", "N/A")
            person_name = extracted_data.get("person_name", "N/A")
            person_address = extracted_data.get("person_address", "N/A")
            country = extracted_data.get("country", "N/A")
            return {"description": description, "project_title": project_title, "person_name": person_name, "person_address": person_address, "country": country}
        except json.JSONDecodeError:
            # Fallback if Gemini doesn't return perfect JSON
            print(f"WARNING: Gemini response not perfect JSON: {response.text}")
            # Try a simpler regex extraction for robustness
            description_match = re.search(r'"project_description":\s*"(.*?)"', response.text, re.DOTALL)
            project_title_match = re.search(r'"project_title":\s*"(.*?)"', response.text)
            person_name_match = re.search(r'"person_name":\s*"(.*?)"', response.text)
            person_address_match = re.search(r'"person_address":\s*"(.*?)"', response.text, re.DOTALL) # DOTALL for multiline address
            country_match = re.search(r'"country":\s*"(.*?)"', response.text)

            description = description_match.group(1).strip() if description_match else response.text.strip() # Use full text as fallback
            project_title = project_title_match.group(1).strip() if project_title_match else "N/A"
            person_name = person_name_match.group(1).strip() if person_name_match else "N/A"
            person_address = person_address_match.group(1).strip() if person_address_match else "N/A"
            country = country_match.group(1).strip() if country_match else "N/A"


            return {"description": description, "project_title": project_title, "person_name": person_name, "person_address": person_address, "country": country}

# Initialize the database and document generator
generator = DocumentGenerator()

# --- Flask Routes ---

@app.route('/api/extract-description', methods=['POST'])
def extract_description():
    print("\n--- Flask: Incoming extract-description request ---")
    print("Request Headers:", request.headers)

    project_file = request.files.get('projectFile')
    gemini_api_key = request.form.get('geminiApiKey') # FormData for this endpoint

    print(f"DEBUG: projectFile received: {'Yes' if project_file else 'No'}")
    print(f"DEBUG: projectFile name: {project_file.filename if project_file else 'N/A'}")
    print(f"DEBUG: geminiApiKey received: {'Yes' if gemini_api_key else 'No'}")
    print(f"DEBUG: geminiApiKey (first 5 chars): '{gemini_api_key[:5] if gemini_api_key else 'N/A'}'")
    print("--- End Flask Debug ---")

    if not project_file:
        print("Flask: Missing projectFile, returning 400")
        return jsonify({"error": "No project file provided."}), 400
    if not gemini_api_key:
        print("Flask: Missing Gemini API Key, returning 400")
        return jsonify({"error": "Gemini API Key not provided."}), 400

    file_content = project_file.read()
    file_type = project_file.mimetype

    try:
        extracted_data = generator.extract_description_with_gemini(file_content, file_type, gemini_api_key)
        print("Flask: Description extracted successfully.")
        return jsonify(extracted_data) # Return the dictionary
    except ValueError as e:
        print(f"Flask: Error extracting description - {str(e)}")
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        print(f"Flask: Unexpected error extracting description - {str(e)}")
        return jsonify({"error": "An internal server error occurred during description extraction."}), 500


@app.route('/api/templates', methods=['POST', 'GET'])
@app.route('/api/templates/<template_name>', methods=['DELETE']) # Specific route for DELETE
def templates_route(template_name=None):
    if request.method == 'POST':
        return upload_template()
    elif request.method == 'GET':
        return get_templates()
    elif request.method == 'DELETE':
        if template_name:
            return delete_template(template_name)
        else:
            return jsonify({"error": "Template name required for deletion"}), 400

def upload_template():
    if 'template' not in request.files:
        return jsonify({"error": "No template file provided"}), 400

    template_file = request.files['template']
    if template_file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if not template_file.filename.endswith('.docx'):
        return jsonify({"error": "Only .docx files are supported for templates"}), 400

    template_name = template_file.filename
    template_content = template_file.read()

    try:
        # Extract format information and the document object for verification
        format_info, _ = generator.extract_template_format(template_content) # _ because we don't need the doc object here

        # Save template content and format_info to database
        if generator.template_db.save_template(template_name, template_content, format_info):
            print(f"Flask: Template '{template_name}' uploaded and processed successfully.")
            return jsonify({"message": "Template uploaded successfully", "name": template_name}), 200
        else:
            print(f"Flask: Failed to save template '{template_name}'.")
            return jsonify({"error": "Failed to save template"}), 500
    except Exception as e:
        print(f"Flask: Error uploading template - {str(e)}")
        return jsonify({"error": f"Failed to process template: {str(e)}"}), 500

def get_templates():
    templates_data = generator.template_db.get_templates()
    templates_list = []
    for name, _, format_info_json in templates_data: # content is not sent to frontend
        try:
            format_info = json.loads(format_info_json)
        except json.JSONDecodeError:
            format_info = {} # Handle cases where format_info might be invalid JSON
            print(f"WARNING: Corrupted format_info for template '{name}'.")
        templates_list.append({"name": name, "format_info": format_info})
    print(f"Flask: Retrieved {len(templates_list)} templates.")
    return jsonify(templates_list)

def delete_template(template_name):
    generator.template_db.delete_template(template_name)
    print(f"Flask: Template '{template_name}' deleted.")
    return jsonify({"message": f"Template '{template_name}' deleted successfully"})

@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    print("\n--- DEBUG: Incoming /api/generate-document request ---")

    content_type = request.headers.get('Content-Type', '')
    print(f"DEBUG: Request Content-Type: {content_type}")

    if 'multipart/form-data' not in content_type:
        print("Flask: Content-Type is not multipart/form-data. Returning 400.")
        return jsonify({"error": "Content-Type must be multipart/form-data"}), 400

    print("DEBUG: Raw request.form data:")
    if request.form:
        for key, value in request.form.items():
            print(f"    {key}: {value}")
    else:
        print("    request.form is empty!")

    print("DEBUG: Raw request.files data:")
    if request.files:
        for key, value in request.files.items():
            print(f"    {key}: {value.filename}")
    else:
        print("    request.files is empty!")

    project_description = request.form.get('projectDescription')
    lead_firm = request.form.get('leadFirm')
    jv_firm = request.form.get('jvFirm')
    document_date_str = request.form.get('date')
    person_name = request.form.get('personName')
    person_address = request.form.get('personAddress')
    project_title = request.form.get('projectTitle') # New
    country = request.form.get('country') # New
    current_location = generator.get_current_location() # Get location from backend (currently hardcoded)


    selected_templates = request.form.getlist('selectedTemplates')

    gemini_api_key = request.form.get('geminiApiKey')

    print(f"DEBUG: Frontend sent 'selectedTemplates': {selected_templates} (Type: {type(selected_templates)})")
    print(f"DEBUG: Number of templates to process: {len(selected_templates)}")

    print(f"DEBUG: projectDescription: {project_description[:50] + '...' if project_description else 'N/A'}")
    print(f"DEBUG: projectTitle: {project_title}") # New
    print(f"DEBUG: leadFirm: {lead_firm}")
    print(f"DEBUG: jvFirm: {jv_firm}")
    print(f"DEBUG: documentDate: {document_date_str}")
    print(f"DEBUG: personName: {person_name}")
    print(f"DEBUG: personAddress: {person_address}")
    print(f"DEBUG: country: {country}") # New
    print(f"DEBUG: current_location: {current_location}")
    print(f"DEBUG: selectedTemplates: {selected_templates}")
    print(f"DEBUG: geminiApiKey (first 5 chars): {gemini_api_key[:5] if gemini_api_key else 'N/A'}")
    print("--- End Flask Debug ---")


    if not all([project_description, lead_firm, selected_templates, gemini_api_key]):
        print("Flask: Missing required fields for document generation, returning 400.")
        return jsonify({"error": "Missing required project details (description, lead firm, selected templates, or Gemini API Key)."}), 400

    document_date_for_gemini = document_date_str if document_date_str else datetime.now().strftime('%Y-%m-%d')


    generated_docs = []
    all_successful = True

    try:
        for i, template_name in enumerate(selected_templates):
            print(f"DEBUG: Processing template {i+1}/{len(selected_templates)}: '{template_name}'")

            templates_in_db = generator.template_db.get_templates()
            template_content = None
            format_info = None

            for name, content, format_info_json in templates_in_db:
                if name == template_name:
                    template_content = content
                    format_info = json.loads(format_info_json)
                    break

            if template_content is None or format_info is None:
                print(f"ERROR: Template '{template_name}' not found in database or format info corrupted.")
                all_successful = False
                continue

            # Generate content using Gemini
            generated_content_lines = generator.generate_content_with_placeholders(
                template_content,
                project_description, lead_firm, jv_firm, document_date_for_gemini, gemini_api_key,
                person_name, current_location, person_address,
                project_title, country # Pass new parameters
            )
            print(f"DEBUG: Gemini generated {len(generated_content_lines)} lines for '{template_name}'.")

            # Apply formatting to the generated content
            final_doc = generator.apply_template_formatting_exact(format_info, generated_content_lines)
            doc_bytes_io = BytesIO()
            final_doc.save(doc_bytes_io)
            doc_bytes_io.seek(0)

            output_filename = f"{os.path.splitext(template_name)[0]}_generated.docx"

            generated_docs.append({
                'name': output_filename,
                'content': doc_bytes_io.getvalue()
            })
            print(f"DEBUG: Successfully added '{output_filename}' to generated_docs.")


    except Exception as e:
        print(f"ERROR: Error generating document(s): {str(e)}")
        return jsonify({"error": f"Error during document generation: {str(e)}"}), 500

    print(f"DEBUG: Final number of documents in generated_docs: {len(generated_docs)}")

    if not generated_docs:
        return jsonify({"error": "No documents could be generated from the selected templates. Please check template names or API key."}), 500

    if len(generated_docs) == 1:
        doc = generated_docs[0]
        return send_file(
            BytesIO(doc['content']),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=doc['name']
        )
    else:
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            for doc in generated_docs:
                zip_file.writestr(doc['name'], doc['content'])

        zip_buffer.seek(0)
        zip_filename = f"generated_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        return send_file(
            zip_buffer,
            mimetype="application/zip",
            as_attachment=True,
            download_name=zip_filename
        )

if __name__ == '__main__':
    generator = DocumentGenerator()
    app.run(debug=True, port=5000)